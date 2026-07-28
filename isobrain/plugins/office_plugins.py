import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook, Workbook
from docx2pdf import convert
from isobrain.core.utils import resolve_path, resolve_target_file

class OfficePlugins:
    # ==================== WORD AUTOMATION NÂNG CAO ====================
    @staticmethod
    def create_word_table(file_name: str = "document.docx", folder_path: str = "", headers_str: str = "", cols_count: int = 0) -> str:
        """Tạo Bảng biểu chuyên nghiệp trong file Word"""
        target_path = resolve_target_file(file_name, folder_path)
        
        # Tự động tạo thư mục cha nếu chưa có
        if not target_path.parent.exists():
            target_path.parent.mkdir(parents=True, exist_ok=True)

        # Mở file cũ hoặc tạo file Word mới
        if target_path.exists():
            doc = Document(target_path)
        else:
            doc = Document()
            doc.add_heading(f"Tài liệu: {target_path.stem}", level=1)

        # Trích xuất danh sách cột từ chuỗi người dùng gõ
        # Ví dụ: "cột 1: stt; cột 2: tên; cột 3: nghĩa; cột 4: ví dụ" -> ['stt', 'tên', 'nghĩa', 'ví dụ']
        raw_headers = re.split(r'[;,]', headers_str)
        cleaned_headers = []
        for h in raw_headers:
            clean_h = re.sub(r'^cột\s*\d+\s*:\s*', '', h.strip(), flags=re.IGNORECASE)
            clean_h = clean_h.strip().strip('"').strip("'")
            if clean_h:
                cleaned_headers.append(clean_h)

        if not cleaned_headers:
            cols = int(cols_count) if cols_count else 3
            cleaned_headers = [f"Cột {i+1}" for i in range(cols)]

        num_cols = len(cleaned_headers)
        
        try:
            # Tạo bảng trong Word
            table = doc.add_table(rows=2, cols=num_cols)
            table.style = 'Table Grid'

            # Điền Header tiêu đề cột
            hdr_cells = table.rows[0].cells
            for i, header_name in enumerate(cleaned_headers):
                hdr_cells[i].text = header_name.upper()
                # Viết hoa in đậm tiêu đề
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            doc.save(target_path)
            folder_str = str(target_path.parent).rstrip("\\")
            return f"[bold green]Thành công![/bold green] Đã tạo Bảng biểu [magenta]{num_cols} cột[/magenta] ({', '.join(cleaned_headers)}) trong file [yellow]{folder_str}\\{target_path.name}[/yellow]."
        except Exception as e:
            return f"[bold red]Lỗi tạo bảng Word:[/bold red] {str(e)}"

    @staticmethod
    def change_word_font(file_path: str, font_name: str) -> str:
        """Đổi font chữ trong Word"""
        path = resolve_target_file(file_path)
        if not path.exists():
            return f"[bold red]Lỗi:[/bold red] File '{path}' không tồn tại!"
        
        try:
            doc = Document(path)
            for p in doc.paragraphs:
                for run in p.runs:
                    run.font.name = font_name
            doc.save(path)
            return f"[bold green]Thành công![/bold green] Đã đổi font file [yellow]{path.name}[/yellow] sang [cyan]{font_name}[/cyan]."
        except PermissionError:
            return f"[bold red]Lỗi Truy Cập:[/bold red] File [yellow]{path.name}[/yellow] đang mở trong Word. Vui lòng đóng file lại!"
        except Exception as e:
            return f"[bold red]Lỗi:[/bold red] {str(e)}"

    @staticmethod
    def replace_text_word(file_path: str, old_text: str, new_text: str) -> str:
        """Tìm kiếm và thay thế văn bản hàng loạt trong Word"""
        path = resolve_target_file(file_path)
        if not path.exists():
            return f"[bold red]Lỗi:[/bold red] File '{path}' không tồn tại!"

        try:
            doc = Document(path)
            replaced_count = 0
            for p in doc.paragraphs:
                if old_text in p.text:
                    p.text = p.text.replace(old_text, new_text)
                    replaced_count += 1
            doc.save(path)
            return f"[bold green]Thành công![/bold green] Đã thay thế [magenta]{replaced_count}[/magenta] đoạn văn chứa từ '[cyan]{old_text}[/cyan]' thành '[cyan]{new_text}[/cyan]'."
        except Exception as e:
            return f"[bold red]Lỗi:[/bold red] {str(e)}"

    @staticmethod
    def convert_word_to_pdf(file_path: str) -> str:
        """Chuyển đổi file Word sang PDF"""
        path = resolve_target_file(file_path)
        if not path.exists():
            return f"[bold red]Lỗi:[/bold red] File '{path}' không tồn tại!"

        pdf_path = path.with_suffix(".pdf")
        try:
            convert(str(path), str(pdf_path))
            return f"[bold green]Thành công![/bold green] Đã xuất file PDF tại: [yellow]{pdf_path}[/yellow]"
        except Exception as e:
            return f"[bold red]Lỗi chuyển PDF:[/bold red] {str(e)}"

    # ==================== EXCEL AUTOMATION NÂNG CAO ====================
    @staticmethod
    def calculate_excel_column(file_path: str, col_letter: str, calc_type: str = "sum") -> str:
        """Tính toán Cột Excel (SUM, AVG, MAX, MIN)"""
        path = resolve_target_file(file_path)
        if not path.exists():
            return f"[bold red]Lỗi:[/bold red] File '{path}' không tồn tại!"

        try:
            wb = load_workbook(path, data_only=True)
            sheet = wb.active
            col_letter = col_letter.upper()

            values = []
            for cell in sheet[col_letter]:
                if isinstance(cell.value, (int, float)):
                    values.append(cell.value)

            if not values:
                return f"[yellow]Không tìm thấy dữ liệu số hợp lệ ở cột {col_letter}.[/yellow]"

            calc_type = calc_type.lower()
            if "trung bình" in calc_type or "avg" in calc_type:
                res = sum(values) / len(values)
                label = "Trung bình"
            elif "lớn nhất" in calc_type or "max" in calc_type:
                res = max(values)
                label = "Giá trị lớn nhất"
            elif "nhỏ nhất" in calc_type or "min" in calc_type:
                res = min(values)
                label = "Giá trị nhỏ nhất"
            else:
                res = sum(values)
                label = "Tổng"

            return f"[bold green]Kết quả:[/bold green] {label} cột [cyan]{col_letter}[/cyan] trong [yellow]{path.name}[/yellow] = [bold magenta]{res:,.2f}[/bold magenta] (từ {len(values)} dòng)."
        except Exception as e:
            return f"[bold red]Lỗi Excel:[/bold red] {str(e)}"