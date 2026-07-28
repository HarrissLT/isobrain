import os
import re
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from isobrain.core.utils import resolve_path

class FilePlugins:
    @staticmethod
    def create_file(file_name: str, folder_path: str) -> str:
        """Tạo 1 file đơn lẻ (.docx, .xlsx, .txt,...)"""
        folder = resolve_path(folder_path)
        clean_file_name = file_name.strip().strip('"').strip("'")
        
        if not folder.exists():
            folder.mkdir(parents=True, exist_ok=True)
            
        full_path = folder / clean_file_name
        
        try:
            ext = full_path.suffix.lower()
            if ext == ".docx":
                doc = Document()
                doc.add_heading(f"Tài liệu: {full_path.stem}", level=1)
                doc.save(full_path)
            elif ext == ".xlsx":
                wb = Workbook()
                wb.save(full_path)
            else:
                with open(full_path, "w", encoding="utf-8") as f:
                    f.write("")
                    
            folder_str = str(folder).rstrip("\\")
            return f"[bold green]Thành công![/bold green] Đã tạo file [cyan]{clean_file_name}[/cyan] tại [yellow]{folder_str}\\{clean_file_name}[/yellow]."
        except Exception as e:
            return f"[bold red]Lỗi tạo file:[/bold red] {str(e)}"

    @staticmethod
    def create_batch_files(raw_names: str, folder_path: str, default_ext: str = "docx") -> str:
        """Tạo HÀNG LOẠT file từ danh sách tên phân cách bởi dấu ; hoặc ,"""
        folder = resolve_path(folder_path)
        if not folder.exists():
            folder.mkdir(parents=True, exist_ok=True)

        delimiters = [";", ",", "\n"]
        pattern = '|'.join(map(re.escape, delimiters))
        raw_list = re.split(pattern, raw_names)

        clean_ext = default_ext.strip().lower().strip(".")
        if clean_ext in ["word", "doc"]:
            clean_ext = "docx"
        elif clean_ext in ["excel", "sheet"]:
            clean_ext = "xlsx"

        created_files = []
        for item in raw_list:
            name = item.strip().strip('"').strip("'")
            if not name:
                continue

            item_path = Path(name)
            ext = item_path.suffix.lower()
            if not ext:
                name = f"{name}.{clean_ext}"
                ext = f".{clean_ext}"

            full_path = folder / name
            try:
                if ext == ".docx":
                    doc = Document()
                    doc.add_heading(f"Tài liệu: {item_path.stem}", level=1)
                    doc.save(full_path)
                elif ext == ".xlsx":
                    wb = Workbook()
                    wb.save(full_path)
                else:
                    with open(full_path, "w", encoding="utf-8") as f:
                        f.write("")
                created_files.append(name)
            except Exception:
                pass

        if not created_files:
            return f"[yellow]Không thể tạo file nào từ danh sách tên '{raw_names}'.[/yellow]"

        file_list_str = "\n".join([f"  • [cyan]{f}[/cyan]" for f in created_files])
        folder_str = str(folder).rstrip("\\")
        return f"[bold green]Thành công![/bold green] Đã tạo [magenta]{len(created_files)}[/magenta] file trong thư mục [yellow]{folder_str}[/yellow]:\n{file_list_str}"

    @staticmethod
    def get_files_by_size(folder_path: str, top_n: int = 5, mode: str = "largest") -> str:
        """
        NÂNG CẤP ĐA CHIỀU: Phân tích file theo kích thước.
        mode = 'largest' (nặng nhất) hoặc 'smallest' (nhẹ nhất)
        """
        folder = resolve_path(folder_path)
        if not folder.exists() or not folder.is_dir():
            return f"[bold red]Lỗi:[/bold red] Thư mục '{folder}' không tồn tại!"

        file_list = []
        for root, _, files in os.walk(folder):
            for f in files:
                p = Path(root) / f
                try:
                    file_list.append((p, p.stat().st_size))
                except Exception:
                    pass

        if not file_list:
            return f"[yellow]Thư mục '{folder}' không có file nào.[/yellow]"

        # Phân loại logic Tăng dần hay Giảm dần
        mode_clean = str(mode).lower()
        is_smallest = "nhẹ" in mode_clean or "nhỏ" in mode_clean or "smallest" in mode_clean
        
        # reverse=False nếu là Nhẹ nhất (bé đến lớn), reverse=True nếu Nặng nhất (lớn đến bé)
        file_list.sort(key=lambda x: x[1], reverse=not is_smallest)
        top_files = file_list[:int(top_n)]

        lines = []
        for p, size in top_files:
            if size >= 1024 ** 3:
                size_str = f"{size / (1024 ** 3):.2f} GB"
            elif size >= 1024 ** 2:
                size_str = f"{size / (1024 ** 2):.2f} MB"
            elif size >= 1024:
                size_str = f"{size / 1024:.2f} KB"
            else:
                size_str = f"{size} Bytes"

            rel_path = p.name if p.parent == folder else p.relative_to(folder)
            lines.append(f"  • [cyan]{rel_path}[/cyan] — [bold magenta]{size_str}[/bold magenta]")

        title_type = "nhẹ nhất" if is_smallest else "nặng nhất"
        folder_str = str(folder).rstrip("\\")
        return f"[bold green]Top {len(top_files)} file {title_type} trong [yellow]{folder_str}[/yellow]:[/bold green]\n" + "\n".join(lines)

    @staticmethod
    def batch_rename(folder_path: str, old_str: str, new_str: str) -> str:
        """Đổi tên hàng loạt file trong thư mục"""
        folder = resolve_path(folder_path)
        if folder.is_file():
            folder = folder.parent
            
        if not folder.exists() or not folder.is_dir():
            return f"[bold red]Lỗi:[/bold red] Thư mục '{folder}' không tồn tại!"

        renamed_count = 0
        try:
            for item in folder.iterdir():
                if item.is_file() and old_str.lower() in item.name.lower():
                    new_name = item.name.replace(old_str, new_str)
                    new_path = item.parent / new_name
                    item.rename(new_path)
                    renamed_count += 1

            if renamed_count == 0:
                return f"[yellow]Không tìm thấy file nào chứa chuỗi '{old_str}' trong thư mục '{folder}'.[/yellow]"

            folder_str = str(folder).rstrip("\\")
            return f"[bold green]Thành công![/bold green] Đã đổi tên [magenta]{renamed_count}[/magenta] file trong [yellow]{folder_str}[/yellow]."
        except Exception as e:
            return f"[bold red]Lỗi đổi tên file:[/bold red] {str(e)}"

    @staticmethod
    def list_files_by_ext(folder_path: str, extension: str) -> str:
        """Liệt kê danh sách file theo định dạng"""
        folder = resolve_path(folder_path)
        if not folder.exists() or not folder.is_dir():
            return f"[bold red]Lỗi:[/bold red] Thư mục '{folder}' không tồn tại!"

        ext = extension.strip().lower()
        if not ext.startswith("."):
            ext = f".{ext}"

        files = [f.name for f in folder.iterdir() if f.is_file() and f.suffix.lower() == ext]
        if not files:
            return f"[yellow]Không tìm thấy file nào có đuôi '{ext}' trong '{folder}'.[/yellow]"

        file_list_str = "\n".join([f"  • [cyan]{name}[/cyan]" for name in files])
        folder_str = str(folder).rstrip("\\")
        return f"[bold green]Tìm thấy {len(files)} file {ext} trong [yellow]{folder_str}[/yellow]:[/bold green]\n{file_list_str}"