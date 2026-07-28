from pathlib import Path
import tempfile
from docx import Document
from docx.shared import Inches
import openpyxl
from isobrain.core.utils import resolve_path

class VizPlugins:
    @staticmethod
    def create_chart_to_word(excel_file: str, word_file: str, col_label: str, col_val: str) -> str:
        """Đọc dữ liệu Excel, vẽ biểu đồ tròn/cột và chèn thẳng vào file Word"""
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            return "[red]Cần cài đặt matplotlib: pip install matplotlib[/red]"

        excel_path = resolve_path(excel_file)
        word_path = resolve_path(word_file)

        if not excel_path.exists():
            return f"[bold red]Lỗi:[/bold red] File Excel '{excel_path}' không tồn tại!"

        try:
            # 1. Đọc dữ liệu từ Excel bằng openpyxl
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            sheet = wb.active
            
            labels = []
            values = []
            
            col_l = col_label.upper()
            col_v = col_val.upper()

            for row in range(2, sheet.max_row + 1):
                lbl = sheet[f"{col_l}{row}"].value
                val = sheet[f"{col_v}{row}"].value
                if lbl is not None and isinstance(val, (int, float)):
                    labels.append(str(lbl))
                    values.append(val)

            if not values:
                return f"[yellow]Không tìm thấy dữ liệu số hợp lệ ở cột {col_v}.[/yellow]"

            # 2. Vẽ biểu đồ bằng Matplotlib
            plt.figure(figsize=(6, 4))
            plt.bar(labels[:10], values[:10], color='skyblue') # Lấy tối đa 10 mục
            plt.title("Biểu Đồ Phân Tích Dữ Liệu Auto-Viz")
            plt.xlabel("Danh mục")
            plt.ylabel("Giá trị")
            plt.xticks(rotation=45)
            plt.tight_layout()

            # Lưu ảnh tạm
            temp_img = Path(tempfile.gettempdir()) / "isobrain_chart.png"
            plt.savefig(temp_img, dpi=150)
            plt.close()

            # 3. Chèn vào file Word
            if word_path.exists():
                doc = Document(word_path)
            else:
                doc = Document()
                doc.add_heading("Báo Cáo Tự Động IsoBrain Auto-Viz", level=1)

            doc.add_paragraph(f"Báo cáo biểu đồ trích xuất từ file {excel_path.name}:")
            doc.add_picture(str(temp_img), width=Inches(5.5))
            doc.save(word_path)

            return f"[bold green]Thành công![/bold green] Đã vẽ biểu đồ từ [yellow]{excel_path.name}[/yellow] và chèn vào file Word: [cyan]{word_path}[/cyan]"

        except Exception as e:
            return f"[bold red]Lỗi Auto-Viz:[/bold red] {str(e)}"